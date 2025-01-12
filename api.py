from flask_restful import Resource, Api, reqparse, fields, marshal_with
from flask import make_response, jsonify
api=Api()
import pandas as pd
from sdv.datasets.local import load_csvs
from sdv.metadata import Metadata
from sdv.multi_table import HMASynthesizer
from sdv.single_table import CTGANSynthesizer
from flask import request, jsonify
import random
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
import io
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from flask_security import auth_required
#================================================output fields=======================================================



#=======================================================================================================
input_parser = reqparse.RequestParser()
input_parser.add_argument('dataset', type=str, required=True, help='dataset is required')
input_parser.add_argument('format', type=str, required=True, help='format is required')
input_parser.add_argument('count', type=int, required=True, help='rows is required')

# Read the CSV file
def generate_employee(num_rows):
    data = load_csvs(folder_name='content/')

    data.keys()

    hotels_table = data['data-bMNd0cfWfQLiuieC0Y0FN']

    # use the head method to inspect the first few rows of the data
    hotels_table.head(3)


    metadata = Metadata.detect_from_dataframes(data)
    metadata.update_column(
        table_name='data-bMNd0cfWfQLiuieC0Y0FN',
        column_name='name',
        sdtype='name'
    )
    metadata.update_column(
        table_name='data-bMNd0cfWfQLiuieC0Y0FN',
        column_name='License',
        sdtype='license_plate'
    )
    metadata.update_column(
        table_name='data-bMNd0cfWfQLiuieC0Y0FN',
        column_name='address',
        sdtype='address',
        pii=True
    )
    metadata.validate()
    metadata.validate_data(data=data)
    # metadata.save_to_json('metadata.json')

    # in the future, you can reload the metadata object from the file
    # metadata = Metadata.load_from_json('metadata.json')


    synthesizer = HMASynthesizer(metadata)
    synthesizer.fit(data)
    synthetic_data = synthesizer.sample(scale=1)
    synthetic_data['data-bMNd0cfWfQLiuieC0Y0FN'].head(3)


    synthesizer = CTGANSynthesizer(metadata, epochs=70)
    synthesizer.fit(synthetic_data['data-bMNd0cfWfQLiuieC0Y0FN'])
    real_data = synthetic_data['data-bMNd0cfWfQLiuieC0Y0FN']
    synthetic_data = synthesizer.sample(num_rows=num_rows)
    synthetic_data.head()
    # synthetic_data.drop(columns=synthetic_data.columns[0], axis=1, inplace=True)
    synthetic_data.to_csv('content/data-bMNd0cfWfQLiuieC0Y0FN.csv', index= False)
    synthetic_data.to_csv('file2.csv', index= False)
    return synthetic_data

def generate_pci(num_rows):
    data = load_csvs(folder_name='content2/')

    data.keys()

    hotels_table = data['findata']

    # use the head method to inspect the first few rows of the data
    hotels_table.head(3)


    metadata = Metadata.detect_from_dataframes(data)
    metadata.update_column(
        table_name='findata',
        column_name='card_holder_name',
        sdtype='name'
    )
    metadata.update_column(
        table_name='findata',
        column_name='expiration_date',
        sdtype='date'
    )
    metadata.update_column(
        table_name='findata',
        column_name='card_provider',
        sdtype='name'
    )
    metadata.validate()
    metadata.validate_data(data=data)
    # metadata.save_to_json('metadata.json')

    # in the future, you can reload the metadata object from the file
    # metadata = Metadata.load_from_json('metadata.json')


    synthesizer = HMASynthesizer(metadata)
    synthesizer.fit(data)
    synthetic_data = synthesizer.sample(scale=1)
    synthetic_data['findata'].head(3)


    synthesizer = CTGANSynthesizer(metadata, epochs=70)
    synthesizer.fit(synthetic_data['findata'])
    real_data = synthetic_data['findata']
    synthetic_data = synthesizer.sample(num_rows=num_rows)
    synthetic_data.head()
    # synthetic_data.drop(columns=synthetic_data.columns[-1], axis=1, inplace=True)
    synthetic_data.to_csv('content2/findata.csv', index= False)
    synthetic_data.to_csv('file2pci.csv', index= False)
    return synthetic_data


def luhn_checksum(card_number):
    def digits_of(n):
        return [int(d) for d in str(n)]
    digits = digits_of(card_number)
    odd_digits = digits[-1::-2]
    even_digits = digits[-2::-2]
    checksum = 0
    checksum += sum(odd_digits)
    for d in even_digits:
        checksum += sum(digits_of(d*2))
    return checksum % 10

def generate_credit_card():
    prefix = random.choice(['4', '5', '37', '6'])
    number = prefix
    while len(number) < 15:
        number += str(random.randint(0, 9))
    check_digit = str((10 - luhn_checksum(int(number + '0'))) % 10)
    return number + check_digit

def generate_expiration_date():
    current_year = datetime.now().year
    year = random.randint(current_year, current_year + 5)
    month = random.randint(1, 12)
    return f"{month:02d}/{str(year)[2:]}"

def generate_cvv():
    return str(random.randint(100, 999))

def generate_card_provider():
    return random.choice(["Visa", "MasterCard", "American Express", "Discover"])

def generate_name():
    first_names = ["John", "Jane", "Michael", "Emily", "David", "Sarah", "Scott", "Jennifer", "Lisa", "Brian"]
    last_names = ["Smith", "Johnson", "Williams", "Brown", "Jones", "Garcia"]
    return f"{random.choice(first_names)} {random.choice(last_names)}"


def generate_pci_data(count):
    data = []
    for _ in range(count):
        data.append({
            "Card holder name": generate_name(),
            "Security code/CVV": generate_cvv(),
            "Expiration date": generate_expiration_date(),
            "Credit Card number": generate_credit_card(),
            "Card provider": generate_card_provider()
        })
    return data

class EmployeeAPI(Resource):
    @auth_required('token')
    def post(self):
        args = input_parser.parse_args()
        dataset = args['dataset']
        format = args['format']
        rows = args['count']

        try:
            if dataset == "Employee":
                df = generate_employee(rows)
            elif dataset == "PCI":
                df = generate_pci(rows)
            else:
                return jsonify({"error": "Invalid dataset specified"}), 400

            if format == "CSV":
                output = df.to_csv(index=False)
                mimetype = 'text/csv'
                filename = f"{dataset}_data.csv"
            elif format == "Excel":
                output = io.BytesIO()
                wb = Workbook()
                ws = wb.active
                ws.title = "Sheet1"

                for r in dataframe_to_rows(df, index=False, header=True):
                    ws.append(r)

                wb.save(output)
                output.seek(0)
                mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                filename = f"{dataset}_data.xlsx"
            elif format == "JSON":
                output = df.to_json(orient='records')
                mimetype = 'application/json'
                filename = f"{dataset}_data.json"
            elif format == "Doc":
                output = io.BytesIO()
                doc = Document()
                doc.add_heading(f'{dataset} Data', 0)

                # Add a table to the document
                table = doc.add_table(rows=1, cols=len(df.columns))
                table.style = 'Table Grid'
                
                # Add headers
                for i, column in enumerate(df.columns):
                    table.cell(0, i).text = column

                # Add data
                for _, row in df.iterrows():
                    cells = table.add_row().cells
                    for i, value in enumerate(row):
                        cells[i].text = str(value)

                doc.save(output)
                output.seek(0)
                mimetype = 'application/msword'
                filename = f"{dataset}_data.doc"
            else:
                return jsonify({"error": "Invalid format specified"}), 400

            response = make_response(output.getvalue() if isinstance(output, io.BytesIO) else output)
            response.headers['Content-Type'] = mimetype
            response.headers['Content-Disposition'] = f'attachment; filename={filename}'
            return response

        except Exception as e:
            return jsonify({"error": str(e)}), 500

api.add_resource(EmployeeAPI, '/employee')



        








        






